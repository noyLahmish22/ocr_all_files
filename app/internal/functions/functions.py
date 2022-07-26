from app.internal.functions.imageprocessing import box, srceen_image
from app.internal import pytesseract
from app.internal.pytesseract import Output
from pdf2image import convert_from_path
import csv
import shutil
import cv2
import docx
from googletrans import Translator
import aspose.words as aw
import os
from app.config.consts import path_dic_data, poppler_path, tessdata_dir_config
from autocorrect import Speller
from app.internal.pytesseract import pytesseract

pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract"


# for docker
# pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

def get_text_from_files(fullpath, cwd, file, lang, test_flag, ex):
    """
    get the file and his language ,process the file and return the text information
    """
    if not test_flag:
        with open(fullpath.encode('utf-8'), "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    text = ""
    if ex == "doc":
        path = (doc_to_doxc(fullpath))
        text = '\n'.join(get_text_from_docx(path))
    elif ex == "docx":
        text = '\n'.join(get_text_from_docx(fullpath))
    elif ex == "pdf":
        text = get_text_from_pdf(fullpath, cwd, lang)
    else:
        get_text_from_image(fullpath, cwd, lang)
        with open(cwd + path_dic_data + 'result_text.txt', encoding='utf-8') as f:
            lines = f.readlines()
        for data in lines:
            text += data
        os.remove(cwd + path_dic_data + '/result_text.txt')
    text = text.replace('-\n', '')
    if lang != 'heb':
        print("hee")
        text = translation(text)
    for file in os.listdir(cwd + path_dic_data):
        os.remove(cwd + path_dic_data + file)

    return text


def get_text_from_docx(fullpath):
    """
    get path of docx file and return the text in it
    """
    doc = docx.Document(fullpath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(translation(para.text))
    return fullText


def doc_to_doxc(file_path):
    """
    convet doc file to docx
    """
    doc = aw.Document(file_path)
    cwd = os.getcwd()
    doc.save(cwd + "/tests/data/Output.docx")
    return cwd + "/tests/data/Output.docx"


def get_text_from_pdf(fullpath, cwd, lang):
    """
    get the path to the pdf file and return the text in it
    """
    pages = convert_from_path(fullpath, 500, poppler_path=poppler_path)
    text = ''
    image_counter = 1
    for page in pages:
        filename = "page_" + str(image_counter) + ".jpg"
        data_path = cwd + path_dic_data + filename
        page.save(data_path, 'JPEG')
        image_counter = image_counter + 1
    filelimit = image_counter - 1
    for i in range(1, filelimit + 1):
        filename = "page_" + str(i) + ".jpg"
        try:
            text += "page:" + filename + "data:" + str(
                (pytesseract.image_to_string(cwd + path_dic_data + filename, lang=str(lang),
                                             config=tessdata_dir_config)))
            os.remove(cwd + path_dic_data + filename)
        except Exception as e:
            print(e)
    return text


def get_text_from_image(fullpath, cwd, lang):
    """
    get fullpath to image and language which response to the image and return the text included in it
    """
    path = srceen_image(fullpath)
    box(path, cwd)
    # details= (pytesseract.image_to_data(fullpath,output_type=Output.DICT,
    #                                     config=tessdata_dir_config, lang=lang))
    details = pytesseract.image_to_data(cwd + path_dic_data + r"\thresh.png", output_type=Output.DICT,
                                        config=tessdata_dir_config, lang=lang)

    total_boxes = len(details['text'])
    threshold_img = cv2.imread(cwd + path_dic_data + r"\thresh.png")

    for sequence_number in range(total_boxes):
        if int(details['conf'][sequence_number]) > 30:
            (x, y, w, h) = (
                details['left'][sequence_number], details['top'][sequence_number],
                details['width'][sequence_number],
                details['height'][sequence_number])
            threshold_img = cv2.rectangle(threshold_img, (x, y), (x + w, y + h), (0, 255, 0), 2)
    cv2.imshow('captured text', threshold_img)
    cv2.waitKey(0)
    cv2.destroyAllWindows()
    parse_text = []
    word_list = []
    last_word = ''
    line_num = 1
    list_lines = {}
    for word in details['text']:
        if word != '':
            if lang == 'eng':
                spelling = Speller(lang='en')
                word = spelling(word)
            word_list.append(word)
            last_word = word
        if (last_word != '' and word == '') or (word == details['text'][-1]):
            parse_text.append(word_list)
            list_lines[line_num] = word_list
            word_list = []
            line_num += 1

    with open(cwd + path_dic_data + '/result_text.txt', 'w', newline="", encoding="utf-8") as file:
        csv.writer(file, delimiter=" ").writerows(parse_text)
    return parse_text


def translation(text):
    """
    get text and translate to Hebrew
    """
    translator = Translator()
    text_trans = translator.translate(text, dest='iw')
    return text_trans.text
